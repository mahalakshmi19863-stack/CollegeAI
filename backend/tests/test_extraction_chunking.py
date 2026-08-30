import os
from pathlib import Path

from docx import Document
from pypdf import PdfWriter
import pytest
from backend.app.rag.chunking import chunker
from backend.app.rag.extraction import clean_text, extractor


def test_text_cleaning():
    raw = "  Hello   world \x00 with  extra \n\n\n\n lines and \u00a0 non-breaking spaces. "
    cleaned = clean_text(raw)
    assert "\x00" not in cleaned
    assert "\u00a0" not in cleaned
    assert "  " not in cleaned
    assert cleaned.startswith("Hello world")


def test_chunking_preserves_page_numbers_and_metadata():
    pages = [
        (1, "The annual tuition fee for B.Tech programs is ₹1,20,000 per year. Semester registration is mandatory."),
        (2, "Hostel fee is ₹50,000 per academic year including standard mess charges."),
    ]

    chunks = chunker.chunk_document_pages(
        pages_content=pages,
        document_id="doc-test-123",
        document_name="Fee Structure 2026",
        document_version=1,
        category="Fees",
        department="General",
    )

    assert len(chunks) == 2
    assert chunks[0].page_number == 1
    assert "₹1,20,000" in chunks[0].content
    assert chunks[0].document_name == "Fee Structure 2026"
    assert chunks[0].category == "Fees"

    assert chunks[1].page_number == 2
    assert "Hostel fee" in chunks[1].content


def test_txt_extraction_cleans_text_and_preserves_single_page(tmp_path: Path):
    file_path = tmp_path / "notice.txt"
    file_path.write_text("  Notice\x00\n\n\nOffice   hours.  ", encoding="utf-8")

    pages, total_pages = extractor.extract(str(file_path), "TXT")

    assert total_pages == 1
    assert pages == [(1, "Notice\n\nOffice hours.")]


def test_docx_extraction_includes_paragraphs_and_tables(tmp_path: Path):
    file_path = tmp_path / "schedule.docx"
    document = Document()
    document.add_paragraph("Academic schedule")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Semester"
    table.cell(0, 1).text = "Spring"
    document.save(file_path)

    pages, total_pages = extractor.extract(str(file_path), "DOCX")

    assert total_pages == 1
    assert pages[0][0] == 1
    assert "Academic schedule" in pages[0][1]
    assert "Semester | Spring" in pages[0][1]


def test_pdf_extraction_preserves_page_numbers_and_total_pages(tmp_path: Path):
    file_path = tmp_path / "empty-pages.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    with file_path.open("wb") as output:
        writer.write(output)

    pages, total_pages = extractor.extract(str(file_path), "PDF")

    assert pages == []
    assert total_pages == 2


def test_extraction_rejects_unsupported_formats(tmp_path: Path):
    file_path = tmp_path / "notes.md"
    file_path.write_text("not supported", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported file format"):
        extractor.extract(str(file_path), "MD")
