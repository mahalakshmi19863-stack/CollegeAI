import logging
import os
import re
from typing import List, Tuple
from pypdf import PdfReader
import docx

logger = logging.getLogger("college_ai.extraction")


def clean_text(text: str) -> str:
    """Normalize whitespace and remove extraction artifacts while preserving meaning."""
    if not text:
        return ""
    # Remove null bytes emitted by some document extractors
    text = text.replace("\x00", "")
    # Replace non-breaking spaces
    text = text.replace("\u00a0", " ")
    # Normalize multiple newlines to at most two
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    # Normalize multiple horizontal spaces to single space
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


class DocumentExtractor:
    @staticmethod
    def extract_from_pdf(file_path: str) -> Tuple[List[Tuple[int, str]], int]:
        """Extract text from a PDF file page by page. Returns list of (page_number, text) and total page count."""
        pages_content: List[Tuple[int, str]] = []
        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            for idx, page in enumerate(reader.pages):
                page_num = idx + 1
                page_text = page.extract_text() or ""
                cleaned = clean_text(page_text)
                if cleaned:
                    pages_content.append((page_num, cleaned))
            return pages_content, total_pages
        except Exception as e:
            logger.error(f"Error extracting PDF text from {file_path}: {e}")
            raise RuntimeError(f"PDF extraction failed: {str(e)}")

    @staticmethod
    def extract_from_docx(file_path: str) -> Tuple[List[Tuple[int, str]], int]:
        """Extract paragraph and table text from a DOCX document."""
        try:
            doc = docx.Document(file_path)
            blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_text = " | ".join(cell for cell in cells if cell)
                    if row_text:
                        blocks.append(row_text)

            full_text = "\n\n".join(blocks)
            cleaned = clean_text(full_text)
            if cleaned:
                return [(1, cleaned)], 1
            return [], 0
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {file_path}: {e}")
            raise RuntimeError(f"DOCX extraction failed: {str(e)}")

    @staticmethod
    def extract_from_txt(file_path: str) -> Tuple[List[Tuple[int, str]], int]:
        """Extract text from a plain TXT file with utf-8 / fallback encoding."""
        try:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="latin-1") as f:
                    content = f.read()

            cleaned = clean_text(content)
            if cleaned:
                return [(1, cleaned)], 1
            return [], 0
        except Exception as e:
            logger.error(f"Error extracting TXT text from {file_path}: {e}")
            raise RuntimeError(f"TXT extraction failed: {str(e)}")

    @classmethod
    def extract(
        cls, file_path: str, file_type: str
    ) -> Tuple[List[Tuple[int, str]], int]:
        """Dispatch extraction based on file format (PDF, DOCX, TXT)."""
        ext = file_type.upper().lstrip(".")
        if ext == "PDF":
            return cls.extract_from_pdf(file_path)
        elif ext == "DOCX":
            return cls.extract_from_docx(file_path)
        elif ext == "TXT":
            return cls.extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_type}")


extractor = DocumentExtractor()
